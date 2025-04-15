import os
from bs4 import BeautifulSoup
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

class HTMLToWordConverter:
    def __init__(self):
        """初始化转换器"""
        self.default_dir = "/Users/cuisijia/Downloads/Hirain_PDF-2/html/html"  # 默认HTML文件夹路径
        
    def read_html_file(self, html_path: str) -> str:
        """
        读取HTML文件内容
        
        Args:
            html_path (str): HTML文件路径
            
        Returns:
            str: HTML文件内容
        """
        try:
            with open(html_path, 'r', encoding='utf-8') as f:
                return f.read()
        except Exception as e:
            print(f"警告：读取HTML文件 {html_path} 时出错: {str(e)}")
            return ""
            
    def process_html_content(self, html_content: str) -> list:
        """
        处理HTML内容，提取需要的内容并保持行结构
        
        Args:
            html_content (str): HTML内容
            
        Returns:
            list: 处理后的内容列表
        """
        try:
            soup = BeautifulSoup(html_content, 'html.parser')
            content_lines = []
            
            # 按照文档顺序处理所有元素
            for element in soup.find_all(['h1', 'table', 'p', 'div', 'span']):
                if element.name == 'h1':
                    # 处理标题
                    text = element.get_text().strip()
                    if text:
                        content_lines.append(text)
                elif element.name == 'table':
                    # 处理表格
                    # 处理表头
                    thead = element.find('thead')
                    if thead:
                        header_text = []
                        for th in thead.find_all('th'):
                            text = th.get_text().strip()
                            if text:
                                header_text.append(text)
                        if header_text:  # 将表头整合成一行
                            content_lines.append('     '.join(header_text))
                    
                    # 处理表格内容
                    tbody = element.find('tbody')
                    if tbody:
                        for tr in tbody.find_all('tr'):
                            row_text = []
                            for td in tr.find_all('td'):
                                text = td.get_text().strip()
                                if text:
                                    row_text.append(text)
                            if row_text:  # 只添加非空行
                                content_lines.append('     '.join(row_text))
                else:
                    # 处理其他文本内容
                    # 检查是否在表格中
                    if not element.find_parent('table'):
                        text = element.get_text().strip()
                        if text and text not in content_lines:  # 避免重复内容
                            content_lines.append(text)
            
            return content_lines
            
        except Exception as e:
            print(f"警告：处理HTML内容时出错: {str(e)}")
            return []
            
    def create_word_document(self, content: list, output_path: str):
        """
        创建Word文档
        
        Args:
            content (list): 文本内容列表
            output_path (str): 输出文件路径
        """
        try:
            # 创建新的Word文档
            doc = Document()
            
            # 设置默认字体和大小
            style = doc.styles['Normal']
            style.font.name = '宋体'
            style.font.size = Pt(12)
            
            # 添加内容
            for line in content:
                paragraph = doc.add_paragraph()
                run = paragraph.add_run(line)
                run.font.name = '宋体'
                run.font.size = Pt(12)
                
                # 设置段落格式
                paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
                paragraph.space_after = Pt(6)  # 段落间距
                
            # 保存文档
            doc.save(output_path)
            print(f"成功创建Word文档: {output_path}")
            
        except Exception as e:
            print(f"创建Word文档时出错: {str(e)}")
            
    def process_directory(self, input_dir: str = None):
        """
        处理指定目录下的所有HTML文件
        
        Args:
            input_dir (str): 输入目录路径，如果为None则使用默认路径
        """
        # 使用默认目录或用户指定的目录
        directory = input_dir if input_dir else self.default_dir
        
        if not os.path.exists(directory):
            print(f"错误：目录不存在: {directory}")
            return
            
        print(f"正在处理目录: {directory}")
        print("发现以下文件:")
        
        # 遍历目录下的所有文件
        for filename in os.listdir(directory):
            if filename.endswith('.html'):
                print(f"- {filename}")
                
                # 构建完整的文件路径
                html_path = os.path.join(directory, filename)
                
                # 读取HTML内容
                html_content = self.read_html_file(html_path)
                
                if html_content:
                    # 处理HTML内容
                    content_lines = self.process_html_content(html_content)
                    
                    if content_lines:
                        # 构建输出文件路径
                        output_filename = os.path.splitext(filename)[0] + '.docx'
                        output_path = os.path.join(directory, output_filename)
                        
                        # 如果docx文件已存在，则删除
                        if os.path.exists(output_path):
                            try:
                                os.remove(output_path)
                                print(f"删除已存在的文件: {output_filename}")
                            except Exception as e:
                                print(f"删除文件 {output_filename} 时出错: {str(e)}")
                        
                        # 创建Word文档
                        self.create_word_document(content_lines, output_path)
                    else:
                        print(f"警告：文件 {filename} 没有可处理的内容")
                else:
                    print(f"警告：无法处理文件 {filename}")

def main():
    converter = HTMLToWordConverter()
    default_dir = "/Users/cuisijia/Downloads/Hirain_PDF-2/html/html"
    
    print("欢迎使用HTML到Word转换工具！")
    print(f"正在处理目录: {default_dir}")
    
    # 处理目录
    converter.process_directory(default_dir)
    print("\n处理完成！")

if __name__ == "__main__":
    main() 